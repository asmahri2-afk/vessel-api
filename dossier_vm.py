"""
dossier_generate.py — Add this to your existing main.py on the Google Cloud VM.

Requires:
    pip install python-docx requests
    (zipfile, io, re are stdlib)

The endpoint downloads each requested .docx template from Cloudflare Pages,
fills all {{placeholders}}, zips the filled files, and returns the zip.
"""

import io
import re
import zipfile
from typing import Any, Dict, List

import requests
from docx import Document
from fastapi import HTTPException
from pydantic import BaseModel

# ── Base URL for templates on Cloudflare Pages ───────────────────────────────
PAGES_BASE = "https://vesseltracker.pages.dev/port-docs"

# ── Template → filename mapping ──────────────────────────────────────────────
TEMPLATE_FILES = {
    "tva-anp":                "tva-anp.docx",
    "tva-marsa":              "tva-marsa.docx",
    "pilotage":               "pilotage.docx",
    "gardiennage":            "gardiennage.docx",
    "timesheet":              "timesheet.docx",
    "manifest-import-entree": "manifest-import-entree.docx",
    "manifest-import-sortie": "manifest-import-sortie.docx",
    "manifest-export-entree": "manifest-export-entree.docx",
    "manifest-export-sortie": "manifest-export-sortie.docx",
    "declaration-import":     "declaration-import.docx",
    "declaration-export":     "declaration-export.docx",
    "overtime":               "overtime.docx",
    "stowaway":               "stowaway.docx",
}


# ── Request model ─────────────────────────────────────────────────────────────
class DossierRequest(BaseModel):
    imo:            str
    port:           str                    # "laayoune" | "dakhla" | "dakhla-anch"
    operation:      str = "import"         # "import" | "export" | "cabotage"
    templates:      List[str]             # list of template IDs to generate
    vessel_name:    str = ""
    flag:           str = ""
    loa:            str = ""
    deadweight:     str = ""
    gross_tonnage:  str = ""
    owner:          str = ""
    cargo:          str = ""
    bl_weight:      str = ""
    shipper:        str = ""
    notify:         str = ""
    from_:          str = ""              # field alias: "from" is a Python keyword
    to:             str = ""
    bc:             str = ""              # Bon de Commande number
    arrival_date:   str = ""
    berthing_date:  str = ""
    departure_date: str = ""
    date:           str = ""              # "date" tag in gardiennage/overtime/pilotage
    today_date:     str = ""
    agent_count:    str = ""
    ste_garde:      str = ""
    expimp:         str = ""              # "Import" or "Export" for overtime
    shift:          str = ""             # e.g. "1er shift, 2ème shift"

    class Config:
        # Allow "from" as a field name via alias
        populate_by_name = True

    @classmethod
    def model_validate(cls, obj, *args, **kwargs):
        # Map "from" key from JSON to "from_" field
        if isinstance(obj, dict) and "from" in obj:
            obj = dict(obj)
            obj["from_"] = obj.pop("from")
        return super().model_validate(obj, *args, **kwargs)


# ── Placeholder replacement (handles split runs in Word) ─────────────────────
def _replace_in_paragraph(para, replacements: Dict[str, str]):
    """Replace {{tag}} in a paragraph, handling runs split by Word's XML."""
    # Reconstruct full text across all runs
    full = "".join(r.text for r in para.runs)
    new  = full
    for tag, val in replacements.items():
        new = new.replace(f"{{{{{tag}}}}}", val)
    if new == full:
        return  # nothing changed — skip to avoid touching formatting
    # Put the replaced text in the first run, blank the rest
    if para.runs:
        para.runs[0].text = new
        for run in para.runs[1:]:
            run.text = ""


def _replace_in_doc(doc: Document, replacements: Dict[str, str]):
    """Replace in body, tables, headers and footers."""
    # Body paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)
    # Tables (all levels of nesting)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)
    # Headers & footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_paragraph(para, replacements)
        for para in section.footer.paragraphs:
            _replace_in_paragraph(para, replacements)


# ── Download template from Cloudflare Pages ───────────────────────────────────
def _download_template(port: str, template_id: str) -> bytes:
    filename = TEMPLATE_FILES.get(template_id)
    if not filename:
        raise HTTPException(status_code=400, detail=f"Unknown template: {template_id}")
    # Port key to folder name (dakhla-anch → dakhla-anch)
    url = f"{PAGES_BASE}/{port}/{filename}"
    try:
        r = requests.get(url, timeout=15)
        r.raise_for_status()
        return r.content
    except requests.RequestException as e:
        raise HTTPException(
            status_code=502,
            detail=f"Failed to fetch template {filename} from {url}: {e}"
        )


# ── Build replacements dict from request ─────────────────────────────────────
def _build_replacements(req: DossierRequest) -> Dict[str, str]:
    return {
        "vessel_name":    req.vessel_name,
        "imo":            req.imo,
        "flag":           req.flag,
        "loa":            req.loa,
        "deadweight":     req.deadweight,
        "gross_tonnage":  req.gross_tonnage,
        "owner":          req.owner,
        "cargo":          req.cargo,
        "bl_weight":      req.bl_weight,
        "shipper":        req.shipper,
        "notify":         req.notify,
        "from":           req.from_,
        "to":             req.to,
        "bc":             req.bc,
        "arrival_date":   req.arrival_date,
        "berthing_date":  req.berthing_date,
        "departure_date": req.departure_date,
        "date":           req.date or req.today_date,
        "today_date":     req.today_date,
        "port":           req.port.replace("-", " ").title(),   # "Laayoune", "Dakhla"
        "agent_count":    req.agent_count,
        "ste_garde":      req.ste_garde,
        "expimp":         req.expimp or req.operation.title(),
        "shift":          req.shift,
    }


# ── FastAPI endpoint ──────────────────────────────────────────────────────────
# Add this to your existing FastAPI app in main.py:
#
#   from dossier_generate import router as dossier_router
#   app.include_router(dossier_router)
#
# OR paste the route function directly into main.py.

from fastapi import APIRouter, Request as FastAPIRequest
from fastapi.responses import Response as FastAPIResponse

router = APIRouter()


@router.post("/dossier/generate")
async def dossier_generate(req: DossierRequest, request: FastAPIRequest):
    """
    Receives JSON payload from Cloudflare Worker.
    Downloads selected .docx templates from Cloudflare Pages,
    fills placeholders, bundles into a zip, returns zip bytes.
    """
    # Verify API secret (same pattern as your existing endpoints)
    api_secret = request.headers.get("X-API-Secret", "")
    import os
    if api_secret != os.environ.get("API_SECRET", ""):
        raise HTTPException(status_code=401, detail="Unauthorized")

    if not req.templates:
        raise HTTPException(status_code=400, detail="No templates selected")
    if not req.port:
        raise HTTPException(status_code=400, detail="Port required")

    replacements = _build_replacements(req)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for tpl_id in req.templates:
            try:
                # Download template
                docx_bytes = _download_template(req.port, tpl_id)

                # Open with python-docx
                doc = Document(io.BytesIO(docx_bytes))

                # Replace all placeholders
                _replace_in_doc(doc, replacements)

                # Save filled docx to memory
                out_buf = io.BytesIO()
                doc.save(out_buf)
                out_buf.seek(0)

                # Add to zip with a clean filename
                fname = TEMPLATE_FILES.get(tpl_id, f"{tpl_id}.docx")
                zf.writestr(fname, out_buf.read())

            except HTTPException:
                raise
            except Exception as e:
                # Log but continue with remaining templates
                print(f"[DOSSIER] Error processing {tpl_id}: {e}")
                continue

    zip_buffer.seek(0)
    zip_bytes = zip_buffer.read()

    if not zip_bytes:
        raise HTTPException(status_code=500, detail="No documents could be generated")

    vessel_slug = re.sub(r"\s+", "_", req.vessel_name.upper()) if req.vessel_name else "VESSEL"
    port_slug   = req.port.upper()
    from datetime import date
    date_slug   = date.today().strftime("%Y%m%d")
    filename    = f"DOSSIER_{vessel_slug}_{port_slug}_{date_slug}.zip"

    return FastAPIResponse(
        content=zip_bytes,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )
