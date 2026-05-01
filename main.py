import json
import time
import logging
import os
import re
import requests
import httpx
import io
from concurrent.futures import ThreadPoolExecutor, as_completed
from copy import copy
from datetime import datetime, timezone
from bs4 import BeautifulSoup
from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel
from typing import Dict, Any, List, Optional
from openpyxl import load_workbook
from openpyxl.styles import Font, Border, Side, PatternFill, Alignment
import zipfile
from docx import Document

# ============================================================
# LOGGING
# ============================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-8s | %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger(__name__)

# ============================================================
# HTTP CONFIG
# ============================================================

import random

_USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36 Edg/120.0.0.0",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
]

def _make_headers(referer: str = "https://www.vesselfinder.com/") -> dict:
    return {
        "User-Agent": random.choice(_USER_AGENTS),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate, br",
        "Referer": referer,
        "DNT": "1",
    }

def _make_mst_headers() -> dict:
    """
    Headers that closely mirror a real Chrome 120 top-level navigation to
    myshiptracking.com.  The Sec-Fetch-* set is what Cloudflare's JS challenge
    checks — without these the TLS fingerprint alone is not enough.
    """
    return {
        "User-Agent": random.choice(_USER_AGENTS),
        # Chrome sends this exact Accept string for document navigations
        "Accept": (
            "text/html,application/xhtml+xml,application/xml;"
            "q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,"
            "application/signed-exchange;v=b3;q=0.7"
        ),
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate, br",
        # Sec-Fetch-* — critical for Cloudflare bot detection
        "Sec-Fetch-Dest":    "document",
        "Sec-Fetch-Mode":    "navigate",
        "Sec-Fetch-Site":    "none",      # direct navigation (no referrer)
        "Sec-Fetch-User":    "?1",        # user-initiated
        "Upgrade-Insecure-Requests": "1",
        "Cache-Control": "max-age=0",
        "DNT": "1",
    }

HEADERS = _make_headers()

MYSHIPTRACKING_URL = "https://www.myshiptracking.com/requests/vesselsonmaptempTTT.php"

API_SECRET         = os.getenv("API_SECRET", "")
EQUASIS_EMAIL      = os.getenv("EQUASIS_EMAIL", "")
EQUASIS_PASSWORD   = os.getenv("EQUASIS_PASSWORD", "")

BATCH_MAX_WORKERS = 2
BATCH_MAX_IMOS    = 50

# ============================================================
# FASTAPI APP + CORS
# ============================================================

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ============================================================
# REQUEST MODELS
# ============================================================

class BatchRequest(BaseModel):
    imos: List[str]

# ============================================================
# UTILITY HELPERS
# ============================================================

def count_decimals(val: Any) -> int:
    if val is None:
        return 0
    s = str(val)
    if "." in s:
        return len(s.split(".")[-1].rstrip("0"))
    return 0

def parse_vf_timestamp(ts: Optional[str]) -> Optional[datetime]:
    if not ts:
        return None
    ts = ts.replace(" UTC", "").strip()
    for fmt in ("%b %d, %Y %H:%M", "%B %d, %Y %H:%M"):
        try:
            return datetime.strptime(ts, fmt).replace(tzinfo=timezone.utc)
        except ValueError:
            continue
    logger.warning(f"Could not parse VF timestamp: '{ts}'")
    return None

def get_vf_age_minutes(last_pos_utc: Optional[str]) -> int:
    dt = parse_vf_timestamp(last_pos_utc)
    if not dt:
        return 999
    age = (datetime.now(timezone.utc) - dt).total_seconds() / 60
    return int(age)

def parse_mst_timestamp(ts: Optional[str]) -> Optional[datetime]:
    """
    Parse MST timestamps which come in ISO-style formats:
      '2024-05-01 14:30', '2024-05-01 14:30:00', '2024-05-01T14:30:00', etc.
    Handles both space- and T-separated variants, with or without seconds.
    """
    if not ts:
        return None
    ts = ts.replace(" UTC", "").replace("Z", "").strip()
    for fmt in (
        "%Y-%m-%dT%H:%M:%S",
        "%Y-%m-%dT%H:%M",
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M",
    ):
        try:
            return datetime.strptime(ts, fmt).replace(tzinfo=timezone.utc)
        except ValueError:
            continue
    logger.warning(f"Could not parse MST timestamp: '{ts}'")
    return None

def get_mst_age_minutes(ts: Optional[str]) -> int:
    dt = parse_mst_timestamp(ts)
    if not dt:
        return 999
    return int((datetime.now(timezone.utc) - dt).total_seconds() / 60)

# ============================================================
# INPUT VALIDATION
# ============================================================

def validate_imo(imo: str) -> bool:
    imo = str(imo).strip()
    if not re.match(r'^\d{7}$', imo):
        return False
    try:
        total = sum(int(imo[i]) * (7 - i) for i in range(6))
        return int(imo[6]) == total % 10
    except Exception:
        return False

# ============================================================
# HTML HELPERS – VESSELFINDER
# ============================================================

def extract_table_data(soup: BeautifulSoup, table_class: str) -> Dict[str, str]:
    data: Dict[str, str] = {}
    tables = soup.find_all(class_=table_class)
    if not tables:
        return data

    for table in tables:
        for row in table.find_all("tr"):
            label_el = row.find(
                class_=lambda x: x and ("tpc1" in x or "tpx1" in x or "n3" in x)
            )
            value_el = row.find(
                class_=lambda x: x and ("tpc2" in x or "tpx2" in x or "v3" in x)
            )
            if not (label_el and value_el):
                continue
            label_parts = [c.strip() for c in label_el.contents if isinstance(c, str)]
            label = " ".join(label_parts).replace(":", "").strip()
            value = value_el.get_text(strip=True)
            if label:
                data[label] = value
    return data

def extract_mmsi(soup: BeautifulSoup, static_data: Dict[str, str]) -> Optional[str]:
    for s in soup.find_all("script"):
        if not s.string:
            continue
        m = re.search(r"MMSI\s*=\s*(\d+)", s.string)
        if m:
            return m.group(1)
    if "MMSI" in static_data:
        v = static_data["MMSI"].strip()
        if v:
            return v
    for key, value in static_data.items():
        if "MMSI" in key.upper():
            v = value.strip()
            if v:
                return v
    return None

# ============================================================
# MYSHIPTRACKING HELPERS
# ============================================================

# Import curl_cffi — required for MST TLS impersonation.
# If missing, install with:  pip install curl_cffi --break-system-packages
try:
    from curl_cffi import requests as curl_requests
    CURL_CFFI_AVAILABLE = True
    logger.info("curl_cffi loaded — MST TLS impersonation enabled")
except ImportError:
    CURL_CFFI_AVAILABLE = False
    logger.warning(
        "curl_cffi not installed — MST scraping will be disabled. "
        "Install with: pip install curl_cffi --break-system-packages"
    )


def _extract_cog_from_scripts(soup: BeautifulSoup) -> Optional[float]:
    """
    Extract Course Over Ground from the hidden canvas_map_generate() JS call.

    MST injects something like:
        canvas_map_generate('canvas_map', 15, 28.123, -12.456, 187.5, 1, 0);
    where the 5th argument is the COG in degrees.

    We try two patterns — one for the exact function name and a broader numeric
    capture — so minor JS changes don't break us.
    """
    for script in soup.find_all("script"):
        raw = script.string
        if not raw or "canvas_map_generate" not in raw:
            continue

        # Pattern A — explicit function call with at least 5 comma-separated args.
        # Skips 4 leading args (any chars) then captures the 5th numeric arg.
        m = re.search(
            r"canvas_map_generate\s*\([^,]+,\s*[^,]+,\s*[^,]+,\s*[^,]+,\s*"
            r"([+-]?\d+(?:\.\d+)?)\s*[,)]",
            raw,
        )
        if m:
            try:
                return float(m.group(1))
            except ValueError:
                pass

        # Pattern B — looser: find "canvas_map_generate" then grab the first
        # 3-digit-ish number that looks like a heading (0–360).
        m2 = re.search(
            r"canvas_map_generate\s*\(.*?(\b[0-2]?\d{1,2}(?:\.\d+)?\b)\s*[,)]",
            raw,
            re.DOTALL,
        )
        if m2:
            try:
                val = float(m2.group(1))
                if 0.0 <= val <= 360.0:
                    return val
            except ValueError:
                pass

    return None


def _parse_mst_port_calls_from_soup(soup: BeautifulSoup) -> List[Dict]:
    """
    Extract the port calls history table from an already-parsed MST vessel page.

    MST renders a <table class="myst-table"> with columns:
      Port | Arrival | Departure | Duration
    Each port cell contains an <a class="pflag"> with an <img title="Country">.
    Date cells contain a <span> whose text is "YYYY-MM-DD HH:MM".

    Returns a list of dicts: {port_name, country, arrived, departed, duration}
    """
    table = soup.find("table", class_="myst-table")
    if not table:
        return []
    tbody = table.find("tbody")
    if not tbody:
        return []

    def _parse_cell_date(cell) -> Optional[str]:
        span = cell.find("span")
        if not span:
            return None
        raw = span.get_text(separator=" ", strip=True)   # "2026-03-22 22:06"
        for fmt in ("%Y-%m-%d %H:%M", "%Y-%m-%d %H:%M:%S"):
            try:
                dt = datetime.strptime(raw, fmt)
                return dt.replace(tzinfo=timezone.utc).isoformat()
            except ValueError:
                continue
        logger.debug(f"MST port calls: could not parse date '{raw}'")
        return None

    results = []
    for row in tbody.find_all("tr"):
        cells = row.find_all("td")
        if len(cells) < 3:
            continue

        a_tag = cells[0].find("a", class_="pflag")
        if not a_tag:
            continue

        port_name = a_tag.get_text(strip=True)
        img       = a_tag.find("img")
        country   = img["title"].strip() if img and img.get("title") else ""

        results.append({
            "port_name": port_name,
            "country":   country,
            "arrived":   _parse_cell_date(cells[1]),
            "departed":  _parse_cell_date(cells[2]),
            "duration":  cells[3].get_text(strip=True) if len(cells) > 3 else "",
        })

    return results


def get_myshiptracking_pos_html(mmsi: str) -> Optional[Dict[str, Any]]:
    """
    Tier 3 — scrape the MST vessel page using curl_cffi with Chrome 120
    TLS impersonation to bypass Cloudflare protection.

    Returns a dict with: lat, lon, sog, cog, last_pos_utc, ais_source, port_calls
    or None on failure.

    NOTE: port_calls is ALWAYS populated from the same HTML fetch — no extra
    HTTP request is ever made.  Callers should pop "port_calls" before using
    the dict as position data.
    """
    if not CURL_CFFI_AVAILABLE:
        logger.warning("curl_cffi not available — skipping MST HTML scrape")
        return None

    url = f"https://www.myshiptracking.com/vessels/mmsi-{mmsi}"

    try:
        response = curl_requests.get(
            url,
            headers=_make_mst_headers(),
            impersonate="chrome120",
            timeout=25,
            verify=True,
            allow_redirects=True,
        )

        if response.status_code == 403:
            logger.warning(
                f"MST HTML returned 403 for MMSI {mmsi} — "
                "Cloudflare challenge not bypassed"
            )
            return None

        if response.status_code != 200:
            logger.warning(f"MST HTML returned HTTP {response.status_code} for MMSI {mmsi}")
            return None

        text = response.text

        # Parse the HTML once — used for both position and port calls.
        soup = BeautifulSoup(text, "html.parser")

        # ------------------------------------------------------------------
        # Primary extraction — canvas_map_generate() JS call
        #
        # MST always injects this call with all position fields in one line:
        #   canvas_map_generate("map_locator", zoom, lat, lon, cog, sog, ...)
        # This is more reliable than scraping the HTML paragraphs which MST
        # restructures frequently.
        # ------------------------------------------------------------------
        map_match = re.search(
            r'canvas_map_generate\s*\(\s*["\'][^"\']*["\']\s*,\s*[\d.]+\s*,'
            r'\s*([+-]?\d+\.?\d*)\s*,\s*([+-]?\d+\.?\d*)\s*,'
            r'\s*([+-]?\d+\.?\d*)\s*,\s*([+-]?\d+\.?\d*)',
            text,
        )
        if map_match:
            lat = float(map_match.group(1))
            lon = float(map_match.group(2))
            cog = float(map_match.group(3))
            sog = float(map_match.group(4))

            # Timestamp lives in the SEO paragraph outside ft-info:
            # "as reported on <strong>2026-04-11 18:17</strong>"
            time_match = re.search(
                r'reported\s+on\b.{0,60}?(\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2})',
                text,
            )
            last_pos_utc = time_match.group(1) if time_match else None

            # Extract port calls from the SAME already-downloaded HTML page.
            port_calls = _parse_mst_port_calls_from_soup(soup)

            logger.info(
                f"MMSI {mmsi} | MST HTML canvas_map: "
                f"lat={lat}, lon={lon}, sog={sog}, cog={cog}, ts={last_pos_utc}, "
                f"port_calls={len(port_calls)}"
            )
            return {
                "lat":          lat,
                "lon":          lon,
                "sog":          sog,
                "cog":          cog,
                "last_pos_utc": last_pos_utc,
                "ais_source":   "myshiptracking_html",
                "port_calls":   port_calls,   # ← piggy-backed, no extra request
            }

        # Position regex failed — still try to return port calls so the page
        # fetch is not wasted (caller can check "port_calls" even on pos fail).
        port_calls = _parse_mst_port_calls_from_soup(soup)
        if port_calls:
            logger.info(
                f"MMSI {mmsi} | MST HTML: position regex failed but "
                f"extracted {len(port_calls)} port calls"
            )
            return {
                "lat": None, "lon": None, "sog": None, "cog": None,
                "last_pos_utc": None,
                "ais_source":   "myshiptracking_html",
                "port_calls":   port_calls,
            }

        logger.warning(f"Could not parse canvas_map_generate from MST HTML for MMSI {mmsi}")
        return None

    except Exception as e:
        logger.warning(f"MST HTML scrape failed for MMSI {mmsi}: {type(e).__name__}: {e}")
        return None


def get_myshiptracking_pos_vessel_api(mmsi: str) -> Optional[Dict[str, Any]]:
    """
    Tier 1 — direct vessel-by-MMSI JSON endpoint.
    Fastest path; also returns a timestamp so age comparison works correctly.
    Uses curl_cffi with Chrome 120 impersonation to pass Cloudflare checks.
    """
    if not CURL_CFFI_AVAILABLE:
        return None

    url = f"https://www.myshiptracking.com/requests/vessel.php?type=json&mmsi={mmsi}"
    headers = {
        "User-Agent":        random.choice(_USER_AGENTS),
        "Accept":            "application/json, text/javascript, */*; q=0.01",
        "Accept-Language":   "en-US,en;q=0.9",
        "Accept-Encoding":   "gzip, deflate, br",
        "Referer":           f"https://www.myshiptracking.com/vessels/mmsi-{mmsi}",
        "X-Requested-With":  "XMLHttpRequest",
    }
    try:
        resp = curl_requests.get(
            url, headers=headers, impersonate="chrome120", timeout=10
        )
        if resp.status_code == 200:
            if not resp.text.strip():
                logger.debug(f"MST vessel API: empty response for MMSI {mmsi} — endpoint inactive")
                return None
            data = resp.json()
            # Normalise field names — MST has used both 'lng' and 'lon' historically
            lat = data.get("lat") or data.get("latitude")
            lon = data.get("lng") or data.get("lon") or data.get("longitude")
            if lat is None or lon is None:
                logger.warning(f"MST vessel API: no lat/lon in response for MMSI {mmsi}")
                return None
            result = {
                "lat":          float(lat),
                "lon":          float(lon),
                "sog":          float(data["speed"])  if data.get("speed")  is not None else None,
                "cog":          float(data["course"]) if data.get("course") is not None else None,
                "last_pos_utc": data.get("received") or data.get("timestamp"),
                "ais_source":   "myshiptracking_api",
            }
            logger.info(
                f"MMSI {mmsi} | MST vessel API: "
                f"lat={result['lat']}, lon={result['lon']}, ts={result['last_pos_utc']}"
            )
            return result
        logger.debug(f"MST vessel API returned HTTP {resp.status_code} for MMSI {mmsi}")
    except Exception as e:
        logger.debug(f"MST vessel API failed for MMSI {mmsi}: {type(e).__name__}: {e}")
    return None


def get_myshiptracking_pos_map_json(
    mmsi: str,
    center_lat: Optional[float],
    center_lon: Optional[float],
    session: requests.Session,
    pad: float = 0.9,
) -> Optional[Dict[str, Any]]:
    """
    Tier 2 — bounding-box map-tile JSON endpoint.
    Requires known VF coordinates to build the bounding box.
    Often returns 403 but kept as a lightweight second attempt.
    Does NOT return a timestamp (map-tile data has none).
    """
    if center_lat is None or center_lon is None:
        return None
    try:
        lat_f, lon_f = float(center_lat), float(center_lon)
    except (TypeError, ValueError):
        return None

    current_year = datetime.now().year
    params = {
        "type":     "json",
        "minlat":   lat_f - pad,  "maxlat": lat_f + pad,
        "minlon":   lon_f - pad,  "maxlon": lon_f + pad,
        "zoom":     15,  "selid": -1,  "seltype": 0,  "timecode": -1,
        "filters":  json.dumps({
            "vtypes": ",0,3,4,6,7,8,9,10,11,12,13", "ports": "1",
            "minsog": 0, "maxsog": 60, "minsz": 0, "maxsz": 500,
            "minyr": 1950, "maxyr": current_year, "status": "",
            "mapflt_from": "", "mapflt_dest": "",
        }),
    }
    try:
        r = session.get(
            MYSHIPTRACKING_URL,
            params=params,
            headers=_make_headers(referer="https://www.myshiptracking.com/"),
            timeout=10,
        )
        if r.status_code != 200:
            logger.warning(f"MST map JSON returned HTTP {r.status_code} for MMSI {mmsi}")
            return None

        lines = [l.strip() for l in r.text.splitlines() if l.strip()]
        if len(lines) < 3:
            return None

        target = str(mmsi).strip()
        for line in lines[2:]:
            parts = line.split("\t") if "\t" in line else line.split()
            if len(parts) >= 7 and parts[2].strip() == target:
                return {
                    "lat":          float(parts[4]),
                    "lon":          float(parts[5]),
                    "sog":          float(parts[6]) if parts[6] else None,
                    "cog":          float(parts[7]) if len(parts) > 7 and parts[7] else None,
                    "last_pos_utc": None,   # map-tile JSON carries no timestamp
                    "ais_source":   "myshiptracking_map",
                }
    except Exception as e:
        logger.warning(f"MST map JSON failed for MMSI {mmsi}: {e}")
    return None

# ============================================================
# MAIN SCRAPER (VF primary + MST 3-tier fallback)
# ============================================================

def scrape_vf_full(imo: str, session: requests.Session) -> Dict[str, Any]:
    url = f"https://www.vesselfinder.com/vessels/details/{imo}"

    r = session.get(url, headers=_make_headers(), timeout=20)

    if r.status_code == 404:
        logger.info(f"IMO {imo} returned 404 from VesselFinder")
        return {"found": False, "imo": imo}
    r.raise_for_status()

    soup = BeautifulSoup(r.text, "html.parser")

    name_el = soup.select_one("h1.title")
    name = name_el.get_text(strip=True) if name_el else f"IMO {imo}"
    dest_el = soup.select_one("div.vi__r1.vi__sbt a._npNa")
    destination = dest_el.get_text(strip=True) if dest_el else ""

    info_icon = soup.select_one("svg.ttt1.info")
    last_pos_utc = info_icon["data-title"] if info_icon and info_icon.has_attr("data-title") else None
    logger.info(f"IMO {imo} | name={name} | last_pos_utc={last_pos_utc}")

    tech_data      = extract_table_data(soup, "tpt1")
    dims_data      = extract_table_data(soup, "tptfix")
    ais_table_data = extract_table_data(soup, "vessel-info-table")
    aparams_data   = extract_table_data(soup, "aparams")
    static_data    = {**tech_data, **dims_data, **ais_table_data, **aparams_data}
    mmsi           = extract_mmsi(soup, static_data)

    draught_val = static_data.get("Current draught") or static_data.get("Draught")
    if not draught_val:
        match = re.search(r"(?:draught|draft)\s+(?:of\s+)?(\d+(?:\.\d+)?)\s*m", soup.get_text(), re.IGNORECASE)
        if match:
            draught_val = f"{match.group(1)} m"

    final_static_data = {
        "imo": imo,
        "vessel_name": name,
        "ship_type": (
            static_data.get("Ship Type")
            or static_data.get("Ship type")
            or static_data.get("Type")
            or ""
        ),
        "flag": (
            soup.select_one("div.title-flag-icon").get("title")
            if soup.select_one("div.title-flag-icon") else None
        ),
        "mmsi": mmsi,
        "draught_m": draught_val or "",
        "deadweight_t":      static_data.get("Deadweight") or static_data.get("DWT"),
        "gross_tonnage":     static_data.get("Gross Tonnage"),
        "year_of_build":     static_data.get("Year of Build"),
        "length_overall_m":  static_data.get("Length Overall"),
        "beam_m":            static_data.get("Beam"),
    }

    vf_lat = vf_lon = sog = cog = None
    djson_div = soup.find("div", id="djson")
    if djson_div and djson_div.has_attr("data-json"):
        try:
            ais = json.loads(djson_div["data-json"])
            vf_lat = float(ais.get("ship_lat")) if ais.get("ship_lat") else None
            vf_lon = float(ais.get("ship_lon")) if ais.get("ship_lon") else None
            sog    = ais.get("ship_sog")
            cog    = ais.get("ship_cog")
            logger.info(f"IMO {imo} | VF AIS: lat={vf_lat}, lon={vf_lon}, sog={sog}, cog={cog}")
        except Exception as e:
            logger.warning(f"IMO {imo} | Failed to parse djson AIS data: {e}")

    # ========== MYSHIPTRACKING FALLBACK — 3 tiers ==========
    mst_data       = None
    mst_port_calls: List[Dict] = []   # port calls piggy-backed from Tier 3 HTML fetch

    if mmsi is not None:
        # Tier 1: direct vessel-by-MMSI API (fastest, has timestamp)
        mst_data = get_myshiptracking_pos_vessel_api(mmsi)

        # Tier 2: bounding-box map-tile JSON (needs VF coords, often 403)
        if not mst_data and vf_lat is not None and vf_lon is not None:
            mst_data = get_myshiptracking_pos_map_json(mmsi, vf_lat, vf_lon, session)

        # Tier 3: curl_cffi full HTML scrape (slowest, most reliable).
        # Also parses port calls from the SAME page — no extra HTTP request needed.
        if not mst_data or mst_data.get("lat") is None:
            html_result = get_myshiptracking_pos_html(mmsi)
            if html_result:
                # Pull out port calls before using the dict as position data
                mst_port_calls = html_result.pop("port_calls", [])
                if html_result.get("lat") is not None:
                    mst_data = html_result

    # Pull the MST timestamp out before the decision so both ages are comparable.
    # We do NOT merge it into last_pos_utc yet — only do that if we actually use MST.
    mst_last_pos_utc: Optional[str] = None
    if mst_data and mst_data.get("last_pos_utc"):
        mst_last_pos_utc = mst_data.pop("last_pos_utc")

    # ========== DECISION LOGIC (VF vs MST) ==========
    #
    # Priority order:
    #   1. VF has no position at all               → always use MST
    #   2. VF is recent (≤30 min) AND not older
    #      than MST                                → prefer VF (fresh signal wins)
    #   3. VF is stale (>60 min)                   → use MST
    #   4. Age difference ≤10 min (effectively
    #      same signal age)                        → tiebreak on coordinate precision
    #   5. MST is meaningfully fresher             → use MST
    #   6. Default                                 → VF
    #
    use_mst = False
    vf_age  = get_vf_age_minutes(last_pos_utc)
    mst_age = get_mst_age_minutes(mst_last_pos_utc) if mst_last_pos_utc else 999

    if mst_data:
        vf_precision  = (count_decimals(vf_lat) + count_decimals(vf_lon)) if vf_lat is not None else 0
        mst_precision = count_decimals(mst_data["lat"]) + count_decimals(mst_data["lon"])

        if vf_lat is None:
            use_mst = True
            logger.info(f"IMO {imo} | Using MST: VF has no position")

        elif abs(vf_age - mst_age) <= 10:
            # Same or close signal age → precision is the tiebreaker
            use_mst = mst_precision > vf_precision
            logger.info(
                f"IMO {imo} | Age tie (vf={vf_age}min, mst={mst_age}min) → "
                f"{'MST' if use_mst else 'VF'} by precision "
                f"(mst={mst_precision} vs vf={vf_precision} decimal places)"
            )

        elif vf_age <= 30 and vf_age < mst_age:
            # VF is recent AND clearly fresher than MST
            use_mst = False
            logger.info(
                f"IMO {imo} | Using VF: recent and fresher "
                f"(vf={vf_age}min, mst={mst_age}min)"
            )

        elif vf_age > 60 and mst_age < vf_age:
            # VF is stale AND MST is fresher
            use_mst = True
            logger.info(
                f"IMO {imo} | Using MST: VF stale and MST fresher "
                f"(vf={vf_age}min, mst={mst_age}min)"
            )

        elif mst_age < vf_age:
            use_mst = True
            logger.info(
                f"IMO {imo} | Using MST: fresher "
                f"(mst={mst_age}min < vf={vf_age}min)"
            )

        else:
            use_mst = False
            logger.info(
                f"IMO {imo} | Using VF: default "
                f"(vf={vf_age}min, mst={mst_age}min)"
            )

    # Only promote the MST timestamp into last_pos_utc when we actually use MST
    if use_mst and mst_last_pos_utc:
        last_pos_utc = mst_last_pos_utc

    if use_mst and mst_data:
        lat, lon   = mst_data["lat"], mst_data["lon"]
        sog        = mst_data.get("sog", sog)
        cog        = mst_data.get("cog", cog)
        ais_source = mst_data.get("ais_source", "myshiptracking")
    else:
        lat, lon   = vf_lat, vf_lon
        ais_source = "vesselfinder"

    logger.info(f"IMO {imo} | Final: lat={lat}, lon={lon}, sog={sog}, source={ais_source}")

    return {
        "found": True,
        "destination": destination,
        "last_pos_utc": last_pos_utc,
        **final_static_data,
        "lat": lat,
        "lon": lon,
        "sog": sog,
        "cog": cog,
        "ais_source": ais_source,
        # Port calls are populated only when Tier 3 HTML scrape fired.
        # Empty list means Tier 1/2 handled position — scrape_vesselfinder.py
        # will request /port-calls/{imo} separately for stale entries.
        "port_calls": mst_port_calls,
    }

# ============================================================
# API ENDPOINTS
# ============================================================

def _check_auth(request: Request, imo: str = ""):
    if API_SECRET:
        client_secret = request.headers.get("X-API-Secret", "")
        if client_secret != API_SECRET:
            logger.warning(f"Unauthorized request for IMO {imo} from {request.client.host}")
            raise HTTPException(status_code=401, detail="Unauthorized")

@app.get("/ping")
def ping():
    return {"ok": True}

@app.get("/vessel-full/{imo}")
def vessel_full(imo: str, request: Request):
    _check_auth(request, imo)

    if not validate_imo(imo):
        logger.warning(f"Invalid IMO rejected: {imo}")
        raise HTTPException(status_code=400, detail="Invalid IMO number")

    with requests.Session() as session:
        try:
            data = scrape_vf_full(imo, session)
        except Exception as e:
            logger.error(f"Scrape failed for IMO {imo}: {e}", exc_info=True)
            raise HTTPException(status_code=502, detail="Upstream scrape failed")

    if not data.get("found"):
        raise HTTPException(status_code=404, detail="Vessel not found")

    return data

@app.get("/equasis/{imo}")
def equasis_endpoint(imo: str, request: Request):
    # 1. Auth Check (Ensure your curl includes -H "X-API-Secret: your_secret")
    _check_auth(request, imo)

    if not validate_imo(imo):
        logger.warning(f"Invalid IMO rejected: {imo}")
        raise HTTPException(status_code=400, detail="Invalid IMO number")

    try:
        # 2. Get the authenticated session
        # Make sure EQUASIS_EMAIL and EQUASIS_PASSWORD are in your .env
        session = _equasis_session()
        
        # 3. Add mandatory headers to bypass simple bot checks
        session.headers.update({
            "Referer": "https://www.equasis.org/EquasisWeb/restricted/Search?fs=Search",
            "Origin": "https://www.equasis.org"
        })

        # 4. Run the scraper logic
        data = _scrape_equasis(imo, session)

        # 5. Check if we actually got data or just the 196-byte redirect
        if not data.get("name") and not data.get("mmsi"):
             logger.error(f"Oracle IP Blocked for IMO {imo} (Empty Data / Redirected)")
             raise HTTPException(
                 status_code=403, 
                 detail="Equasis blocked the request (Oracle Cloud IP range). Use a proxy or check session."
             )
        
        return data

    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Equasis Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/port-calls/{imo}")
def port_calls_endpoint(imo: str, request: Request, mmsi: str = ""):
    """
    Fetch port calls for a vessel from the MST HTML page.

    Called by scrape_vesselfinder.py post-loop when port calls are stale AND
    Tier 3 did not fire during the main vessel-full scrape (i.e. Tier 1 or 2
    handled position, so no HTML page was fetched yet).

    Query param:
        mmsi  — required, 9-digit MMSI of the vessel

    Returns: { imo, count, port_calls: [...] }
    """
    _check_auth(request, imo)

    if not validate_imo(imo):
        raise HTTPException(status_code=400, detail="Invalid IMO number")

    if not mmsi or not re.match(r'^\d{7,9}$', mmsi):
        raise HTTPException(status_code=400, detail="mmsi query param required (7-9 digits)")

    if not CURL_CFFI_AVAILABLE:
        raise HTTPException(status_code=503, detail="curl_cffi not available on this instance")

    result = get_myshiptracking_pos_html(mmsi)

    if result is None:
        logger.warning(f"port-calls/{imo}: MST HTML fetch returned None for MMSI {mmsi}")
        return {"imo": imo, "count": 0, "port_calls": []}

    calls = result.get("port_calls", [])
    logger.info(f"port-calls/{imo}: returning {len(calls)} calls for MMSI {mmsi}")
    return {"imo": imo, "count": len(calls), "port_calls": calls}


@app.post("/vessel-batch")
def vessel_batch(body: BatchRequest, request: Request):
    _check_auth(request)

    imos = list(dict.fromkeys(body.imos))

    if len(imos) > BATCH_MAX_IMOS:
        raise HTTPException(
            status_code=400,
            detail=f"Too many IMOs — max {BATCH_MAX_IMOS} per batch",
        )

    invalid = [imo for imo in imos if not validate_imo(imo)]
    if invalid:
        raise HTTPException(status_code=400, detail=f"Invalid IMOs: {', '.join(invalid)}")

    results: Dict[str, Any] = {}
    errors:  Dict[str, str] = {}

    def fetch_one(imo: str) -> tuple:
        try:
            time.sleep(random.uniform(2, 5))
            with requests.Session() as session:
                data = scrape_vf_full(imo, session)
            if not data.get("found"):
                return imo, None, "Vessel not found"
            return imo, data, None
        except Exception as e:
            logger.error(f"Batch scrape failed for IMO {imo}: {e}")
            return imo, None, str(e)

    logger.info(f"Batch request: {len(imos)} vessels, {BATCH_MAX_WORKERS} workers")

    with ThreadPoolExecutor(max_workers=BATCH_MAX_WORKERS) as executor:
        futures = {executor.submit(fetch_one, imo): imo for imo in imos}
        for future in as_completed(futures):
            imo, data, error = future.result()
            if error:
                errors[imo] = error
            else:
                results[imo] = data

    logger.info(f"Batch complete: {len(results)} success, {len(errors)} errors")

    return {
        "results": results,
        "errors":  errors,
        "total":   len(imos),
        "success": len(results),
        "failed":  len(errors),
    }

# ============================================================
# EQUASIS SCRAPER
# ============================================================
EQUASIS_HOMEPAGE_URL   = "https://www.equasis.org/EquasisWeb/public/HomePage"
EQUASIS_LOGIN_POST_URL = "https://www.equasis.org/EquasisWeb/authen/HomePage?fs=HomePage"
EQUASIS_VESSEL_URL     = "https://www.equasis.org/EquasisWeb/restricted/ShipInfo"

def _equasis_session() -> requests.Session:
    if not EQUASIS_EMAIL or not EQUASIS_PASSWORD:
        raise HTTPException(status_code=503, detail="Equasis credentials not configured")
    session = requests.Session()
    session.headers.update(_make_headers(referer="https://www.equasis.org/"))

    # Step 1: GET the public homepage to establish JSESSIONID + grab CSRF token
    login_page = session.get(EQUASIS_HOMEPAGE_URL, timeout=15)
    login_page.raise_for_status()
    login_soup = BeautifulSoup(login_page.text, "html.parser")
    token_input = login_soup.find("input", {"name": "j_token"})
    token = token_input["value"] if token_input else ""

    # Step 2: POST credentials to the authen endpoint
    login_data = {
        "j_token":    token,
        "j_email":    EQUASIS_EMAIL,
        "j_password": EQUASIS_PASSWORD,
        "submit":     "Login",
    }
    login_resp = session.post(EQUASIS_LOGIN_POST_URL, data=login_data, timeout=15)
    login_resp.raise_for_status()
    if "logout" not in login_resp.text.lower() and "j_password" in login_resp.text.lower():
        raise HTTPException(status_code=401, detail="Equasis login failed — check credentials")
    return session

def _scrape_equasis(imo: str, session: requests.Session) -> Dict[str, Any]:
    params = {"P_IMO": imo}
    # Equasis often expects a specific Referer to allow the search
    headers = {
        "Referer": "https://www.equasis.org/EquasisWeb/restricted/Search?fs=Search"
    }
    
    resp = session.get(EQUASIS_VESSEL_URL, params=params, headers=headers, timeout=20)
    resp.raise_for_status()

    # --- BLOCK/REDIRECT DETECTION ---
    # If content is tiny or contains "Login", the session is dead or Oracle IP is blocked.
    if len(resp.text) < 1000 or "Login" in resp.text or "Restricted" in resp.text:
        # We raise a custom error so your loop can catch it and re-login
        raise ConnectionRefusedError(f"Equasis session expired or IP blocked (Response size: {len(resp.text)}b)")

    soup = BeautifulSoup(resp.text, "html.parser")

    # ─────────────────────────────────────────────────────────────────────
    # Vessel name — Robust extraction
    # ─────────────────────────────────────────────────────────────────────
    name = ""
    # Primary: Find in H4
    for h4 in soup.find_all("h4"):
        txt = h4.get_text(strip=True)
        if "IMO" in txt and imo in txt:
            name = re.split(r"\s*IMO", txt, flags=re.I)[0].strip().rstrip("-").strip()
            break
    
    # Secondary: Fallback to page title if H4 parsing fails
    if not name:
        title = soup.title.string if soup.title else ""
        if title and "Equasis" in title and "-" in title:
            name = title.split("-")[0].strip()

    # ─────────────────────────────────────────────────────────────────────
    # Bootstrap-grid key/value fields
    # ─────────────────────────────────────────────────────────────────────
    info: Dict[str, str] = {}
    for b in soup.find_all("b"):
        label = b.get_text(strip=True).rstrip(":")
        if not label: continue
        
        # Using a more direct col-sibling approach for Bootstrap layouts
        label_col = b.find_parent("div", class_=re.compile(r"col-"))
        if not label_col: continue
        
        # Try to find the immediate next sibling div (the value column)
        value_col = label_col.find_next_sibling("div", class_=re.compile(r"col-"))
        if not value_col: continue

        val_text = value_col.get_text(strip=True)
        
        # Flag logic (Icon detection)
        if not val_text:
            img = value_col.find("img")
            if img and img.get("src"):
                m = re.search(r"/([A-Z]{2,3})\.(?:png|gif|jpg)", img["src"], re.I)
                if m: val_text = m.group(1).upper()

        # Handle Flag Country (usually in the 3rd column over in the row)
        if label.lower().startswith("flag"):
            country_col = value_col.find_next_sibling("div", class_=re.compile(r"col-"))
            if country_col:
                country = re.sub(r"^\(|\)$", "", country_col.get_text(strip=True)).strip()
                if country:
                    val_text = f"{val_text} ({country})" if val_text else country

        if val_text:
            info[label] = val_text

    # ─────────────────────────────────────────────────────────────────────
    # Companies table — find row with role = "Registered owner"
    # ─────────────────────────────────────────────────────────────────────
    equasis_owner = ""
    equasis_address = ""
    for table in soup.find_all("table"):
        rows = table.find_all("tr")
        if not rows: continue
        
        headers = [th.get_text(strip=True).lower() for th in rows[0].find_all(["th", "td"])]
        if not any(h in " ".join(headers) for h in ["role", "address"]): continue

        try:
            i_role = next(i for i, h in enumerate(headers) if "role" in h)
            i_name = next(i for i, h in enumerate(headers) if "name" in h and "compan" in h)
            i_address = next(i for i, h in enumerate(headers) if "address" in h)
        except StopIteration: continue

        for tr in rows[1:]:
            cells = tr.find_all(["td", "th"])
            if len(cells) <= max(i_role, i_name, i_address): continue
            
            role = cells[i_role].get_text(strip=True).lower()
            if "registered owner" in role:
                equasis_owner = cells[i_name].get_text(strip=True)
                equasis_address = cells[i_address].get_text(strip=True)
                break
        if equasis_owner: break

    # ─────────────────────────────────────────────────────────────────────
    # Final Result Construction
    # ─────────────────────────────────────────────────────────────────────
    return {
        "imo":             imo,
        "name":            name,
        "flag":            info.get("Flag") or info.get("Flag State"),
        "gross_tonnage":   info.get("Gross tonnage") or info.get("Gross Tonnage") or info.get("GT"),
        "deadweight_t":    info.get("Deadweight") or info.get("DWT"),
        "ship_type":       info.get("Type of ship") or info.get("Ship type") or info.get("Vessel type"),
        "year_of_build":   info.get("Year of build") or info.get("Year Built"),
        "call_sign":       info.get("Call Sign") or info.get("Callsign"),
        "mmsi":            info.get("MMSI"),
        "equasis_owner":   equasis_owner,
        "equasis_address": equasis_address,
        "class_society":   info.get("Classification society") or "",
        "pi_club":         info.get("P&I Club") or info.get("P&I club") or "",
        "raw_info":        info,
    }

# ============================================================
# SOF GENERATOR
# ============================================================

DAYS = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]

class SOFRow(BaseModel):
    date:    Optional[str] = ''
    wfrom:   Optional[str] = ''
    wto:     Optional[str] = ''
    sfrom:   Optional[str] = ''
    sto:     Optional[str] = ''
    cranes:  Optional[str] = ''
    qty:     Optional[str] = ''
    remarks: Optional[str] = ''

class SOFData(BaseModel):
    agent:             Optional[str] = ''
    vessel:            Optional[str] = ''
    port:              Optional[str] = ''
    owners:            Optional[str] = ''
    cargo:             Optional[str] = ''
    bl_weight:         Optional[str] = ''
    bl_number:         Optional[str] = ''
    operation_type:    Optional[str] = ''
    nor_accepted:      Optional[str] = ''
    port_hours:        Optional[str] = ''
    general_remarks:   Optional[str] = ''
    remarks:           Optional[str] = ''
    master_remarks:    Optional[str] = ''
    berthed_date:      Optional[str] = ''
    berthed_time:      Optional[str] = ''
    disch_start_date:  Optional[str] = ''
    disch_start_time:  Optional[str] = ''
    disch_end_date:    Optional[str] = ''
    disch_end_time:    Optional[str] = ''
    cargo_docs_date:   Optional[str] = ''
    cargo_docs_time:   Optional[str] = ''
    sailing_date:      Optional[str] = ''
    sailing_time:      Optional[str] = ''
    eosp_date:         Optional[str] = ''
    eosp_time:         Optional[str] = ''
    nor_tender_date:   Optional[str] = ''
    nor_tender_time:   Optional[str] = ''
    anchor_drop_date:  Optional[str] = ''
    anchor_drop_time:  Optional[str] = ''
    anchor_weigh_date: Optional[str] = ''
    anchor_weigh_time: Optional[str] = ''
    pilot_date:        Optional[str] = ''
    pilot_time:        Optional[str] = ''
    rows:              List[SOFRow] = []

def fmt_dt(date: str, time: str) -> str:
    if not date:
        return ''
    try:
        parts = date.split('-')
        d = f"{parts[2]}/{parts[1]}/{parts[0]}" if len(parts) == 3 else date
    except Exception:
        d = date
    t = time.replace(':', '') if time else ''
    return f"{d} at   {t} hr" if t else d

def get_day_name(date_str: str) -> str:
    if not date_str:
        return ''
    try:
        dt = datetime.strptime(date_str, '%Y-%m-%d')
        return DAYS[dt.weekday()]
    except Exception:
        return ''

SOF_TEMPLATE_BYTES: Optional[bytes] = None

async def get_sof_template() -> bytes:
    global SOF_TEMPLATE_BYTES
    if SOF_TEMPLATE_BYTES:
        return SOF_TEMPLATE_BYTES
    url = 'https://asmahri2-afk.github.io/test/SOF_TEMPLATE.xlsx'
    async with httpx.AsyncClient(timeout=15) as client:
        r = await client.get(url)
        r.raise_for_status()
        SOF_TEMPLATE_BYTES = r.content
        logger.info(f"SOF template loaded: {len(SOF_TEMPLATE_BYTES)} bytes")
        return SOF_TEMPLATE_BYTES

@app.post('/sof/generate')
async def sof_generate(data: SOFData, request: Request):
    if API_SECRET:
        client_secret = request.headers.get('X-API-Secret', '')
        if client_secret != API_SECRET:
            raise HTTPException(status_code=401, detail='Unauthorized')

    try:
        template_bytes = await get_sof_template()
        wb = load_workbook(io.BytesIO(template_bytes))
        ws = wb.active

        operation_verb = 'loading' if (data.operation_type or '').lower() == 'export' else 'discharging'

        tag_values = {
            '{{AGENT}}':          data.agent or '',
            '{{VESSEL_NAME}}':    data.vessel or '',
            '{{PORT}}':           data.port or '',
            '{{OWNERS}}':         data.owners or '',
            '{{CARGO}}':          data.cargo or '',
            '{{BL_WEIGHT}}':      data.bl_weight or '',
            '{{BL_NUMBER}}':      data.bl_number or '',
            '{{PORT_HOURS}}':     data.port_hours or '',
            '{{GENERAL_REMARKS}}': data.general_remarks or '',
            '{{REMARKS}}':        data.remarks or '',
            '{{MASTER_REMARKS}}': data.master_remarks or '',
            '{{NOR_ACCEPTED}}':   data.nor_accepted or '',
            '{{OPERATION_VERB}}': operation_verb,
            '{{BERTHED_DATE}} at {{BERTHED_TIME}} hr':           fmt_dt(data.berthed_date,      data.berthed_time),
            '{{DISCH_START_DATE}} at {{DISCH_START_TIME}} hr':   fmt_dt(data.disch_start_date,  data.disch_start_time),
            '{{DISCH_END_DATE}} at {{DISCH_END_TIME}} hr':       fmt_dt(data.disch_end_date,    data.disch_end_time),
            '{{CARGO_DOCS_DATE}} at {{CARGO_DOCS_TIME}} hr':     fmt_dt(data.cargo_docs_date,   data.cargo_docs_time),
            '{{SAILING_DATE}} at {{SAILING_TIME}} hr':           fmt_dt(data.sailing_date,      data.sailing_time),
            '{{EOSP_DATE}} at {{EOSP_TIME}} hr':                 fmt_dt(data.eosp_date,         data.eosp_time),
            '{{NOR_TENDER_DATE}} at {{NOR_TENDER_TIME}} hr':     fmt_dt(data.nor_tender_date,   data.nor_tender_time),
            '{{ANCHOR_DROP_DATE}} at {{ANCHOR_DROP_TIME}} hr':  fmt_dt(data.anchor_drop_date,  data.anchor_drop_time),
            '{{ANCHOR_WEIGH_DATE}} at {{ANCHOR_WEIGH_TIME}} hr': fmt_dt(data.anchor_weigh_date, data.anchor_weigh_time),
            '{{PILOT_DATE}} at {{PILOT_TIME}} hr':               fmt_dt(data.pilot_date,        data.pilot_time),
            'M/V {{VESSEL_NAME}}': f"M/V {data.vessel or ''}",
        }

        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    val = cell.value
                    for tag, replacement in tag_values.items():
                        if tag in val:
                            val = val.replace(tag, replacement)
                    cell.value = val

        marker_row = template_row = end_row = None
        for row in ws.iter_rows():
            for cell in row:
                if cell.value == '{{#EACH_ROW}}':  marker_row   = cell.row
                elif cell.value == '{{ROW_DATE}}':  template_row = cell.row
                elif cell.value == '{{/EACH_ROW}}': end_row      = cell.row

        if marker_row and template_row and data.rows:
            tpl_styles = {}
            for col in range(1, 12):
                c = ws.cell(row=template_row, column=col)
                tpl_styles[col] = {
                    'font':      copy(c.font),
                    'border':    copy(c.border),
                    'fill':      copy(c.fill),
                    'alignment': copy(c.alignment),
                }

            for r in filter(None, [marker_row, template_row, end_row]):
                for col in range(1, 12):
                    ws.cell(row=r, column=col).value = None

            # Unmerge any merged cells in the ops-log writing zone
            write_end = marker_row + len(data.rows) + 5
            merges_to_remove = [
                str(rng) for rng in ws.merged_cells.ranges
                if rng.min_row >= marker_row and rng.max_row <= write_end
            ]
            for mr in merges_to_remove:
                ws.unmerge_cells(mr)

            for i, row_data in enumerate(data.rows):
                r = marker_row + i
                values = [
                    fmt_dt(row_data.date, '').split(' ')[0] if row_data.date else '',
                    get_day_name(row_data.date),
                    row_data.wfrom.replace(':', '') if row_data.wfrom else '',
                    row_data.wto.replace(':', '')   if row_data.wto   else '',
                    row_data.sfrom.replace(':', '') if row_data.sfrom else '',
                    row_data.sto.replace(':', '')   if row_data.sto   else '',
                    row_data.cranes  or '',
                    row_data.qty     or '',
                    '',
                    row_data.remarks or '',
                    '',
                ]
                for col, val in enumerate(values, 1):
                    cell = ws.cell(row=r, column=col)
                    cell.value = val if val else None
                    s = tpl_styles.get(col, {})
                    if s.get('font'):      cell.font      = s['font']
                    if s.get('border'):    cell.border    = s['border']
                    if s.get('fill'):      cell.fill      = s['fill']
                    if s.get('alignment'): cell.alignment = s['alignment']

        # ── COMANAV logo injection ─────────────────────────────────────────────
        if (data.agent or '').upper() == 'COMANAV' and ws._images:
            try:
                comanav_url = 'https://asmahri2-afk.github.io/test/logo-comanav.png'
                async with httpx.AsyncClient(timeout=10) as client:
                    logo_resp = await client.get(comanav_url)
                    logo_resp.raise_for_status()
                    logo_bytes = logo_resp.content

                old_img = ws._images[0]
                old_img.ref = io.BytesIO(logo_bytes)
                logger.info(f"Swapped logo to COMANAV ({len(logo_bytes)} bytes)")
            except Exception as e:
                logger.warning(
                    f"Logo swap failed (non-critical): {type(e).__name__}: {e}",
                    exc_info=True,
                )

        # ── Border repair for ops-log J/K columns ──────────────────────────
        # openpyxl can corrupt the J (col 10) / K (col 11) borders on rows
        # whose internal style index differs from the majority (e.g. rows 34,
        # 43 in the current template).  The right-medium border migrates from
        # K onto J, making the vertical line "move inside" visually.
        # Fix: enforce the correct border on every ops-log row (29-58).
        _border_J = Border(left=Side(style='thin'))
        _border_K = Border(right=Side(style='medium'))
        for _r in range(29, 59):
            ws.cell(row=_r, column=10).border = _border_J
            ws.cell(row=_r, column=11).border = _border_K
        # Remove any spurious J:K merges in the ops-log area
        _spurious = [
            str(rng) for rng in ws.merged_cells.ranges
            if rng.min_col == 10 and rng.max_col == 11
            and 29 <= rng.min_row <= 58
        ]
        for _m in _spurious:
            ws.unmerge_cells(_m)

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)

        vessel   = (data.vessel or 'VESSEL').replace(' ', '_')
        port     = (data.port   or 'PORT').replace(' ', '_')
        date_str = datetime.now().strftime('%Y%m%d')
        filename = f"SOF_{vessel}_{port}_{date_str}.xlsx"

        return Response(
            content=buf.read(),
            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'},
        )

    except Exception as e:
        logger.error(f"SOF generation failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"SOF generation failed: {str(e)}")


# ============================================================
# DOSSIER GENERATOR
# ============================================================

DOSSIER_PAGES_BASE = "https://vesseltracker.pages.dev/port-docs"

DOSSIER_TEMPLATE_FILES = {
    "tva-anp":                "tva-anp.docx",
    "tva-marsa":              "tva-marsa.docx",
    "pilotage":               "pilotage.docx",
    "gardiennage":            "gardiennage.docx",
    "timesheet":              "timesheet.docx",
    "manifest-import-entree": "manifest-import-entree.docx",
    "manifest-import-sortie": "manifest-import-sortie.docx",
    "manifest-export-entree": "manifest-export-entree.docx",
    "manifest-export-sortie": "manifest-export-sortie.docx",
    "declaration-import":     "declaration-import.docx",
    "declaration-export":     "declaration-export.docx",
    "overtime":               "overtime.docx",
    "stowaway":               "stowaway.docx",
}

class CargoItem(BaseModel):
    description: str
    weight: str

class DossierRequest(BaseModel):
    imo:            str
    port:           str
    operation:      Optional[str] = "import"
    templates:      List[str]
    vessel_name:    Optional[str] = ""
    flag:           Optional[str] = ""
    loa:            Optional[str] = ""
    deadweight:     Optional[str] = ""
    gross_tonnage:  Optional[str] = ""
    owner:          Optional[str] = ""
    cargo:          Optional[str] = ""           # kept for compatibility
    bl_weight:      Optional[str] = ""           # kept for compatibility
    cargo_items:    Optional[List[CargoItem]] = []   # new multi-cargo field
    shipper:        Optional[str] = ""
    notify:         Optional[str] = ""
    from_port:      Optional[str] = ""   # mapped from "from" key in JSON
    to_port:        Optional[str] = ""   # mapped from "to" key in JSON
    bc:             Optional[str] = ""
    arrival_date:   Optional[str] = ""
    berthing_date:  Optional[str] = ""
    departure_date: Optional[str] = ""
    date:           Optional[str] = ""
    today_date:     Optional[str] = ""
    agent_count:    Optional[str] = ""
    ste_garde:      Optional[str] = ""
    expimp:         Optional[str] = ""
    shift:          Optional[str] = ""

    @classmethod
    def model_validate(cls, obj, *args, **kwargs):
        """Map JSON keys 'from' and 'to' to 'from_port' and 'to_port'."""
        if isinstance(obj, dict):
            obj = dict(obj)
            if "from" in obj:
                obj["from_port"] = obj.pop("from")
            if "to" in obj and "to_port" not in obj:
                obj["to_port"] = obj.pop("to")
        return super().model_validate(obj, *args, **kwargs)


# ============================================================
# DOSSIER GENERATOR – FIXED FUNCTIONS (multi-cargo + pattern replacement)
# ============================================================

def _dossier_replace_paragraph(para, replacements):
    """
    Replace {{tag}} placeholders in a paragraph, preserving <w:br/> (line
    breaks) and <w:tab/> (tabs). Converts `\n` to <w:br/>.

    Special handling: if the paragraph contains the exact pattern
    "{{cargo}} : {{bl_weight}}" (with optional spaces), we replace it with
    a single {{cargo}} placeholder whose value is a newline‑separated list of
    "description : weight" pairs, and clear the {{bl_weight}} placeholder.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not para.runs:
        return

    # ---- Detect and pre‑process the special pattern ----
    full_text = "".join(run.text for run in para.runs)
    pattern = r'\{\{cargo\}\}\s*:\s*\{\{bl_weight\}\}'
    if re.search(pattern, full_text):
        # Build combined lines from cargo and bl_weight
        cargo_str = replacements.get("cargo", "")
        bl_str = replacements.get("bl_weight", "")
        cargo_lines = cargo_str.split('\n') if cargo_str else []
        bl_lines = bl_str.split('\n') if bl_str else []
        max_len = max(len(cargo_lines), len(bl_lines))
        combined_lines = []
        for i in range(max_len):
            desc = cargo_lines[i] if i < len(cargo_lines) else ""
            weight = bl_lines[i] if i < len(bl_lines) else ""
            if desc or weight:
                combined_lines.append(f"{desc}{' : ' + weight if weight else ''}".strip())
        combined = "\n".join(combined_lines)
        # Override replacements: set cargo to combined, clear bl_weight
        replacements = dict(replacements)
        replacements["cargo"] = combined
        replacements["bl_weight"] = ""
        # Replace the pattern in the paragraph text with {{cargo}} (so only one placeholder remains)
        for run in para.runs:
            if run.text:
                run.text = re.sub(pattern, '{{cargo}}', run.text)

    # ---- Standard placeholder replacement (newline → line break) ----
    BR = "\uE000"
    TAB = "\uE001"

    W_T   = qn('w:t')
    W_BR  = qn('w:br')
    W_TAB = qn('w:tab')
    W_CR  = qn('w:cr')

    parts = []
    for run in para.runs:
        for child in run._element:
            if child.tag == W_T:
                parts.append(child.text or '')
            elif child.tag in (W_BR, W_CR):
                parts.append(BR)
            elif child.tag == W_TAB:
                parts.append(TAB)
    full = "".join(parts)

    new = full
    for tag, val in replacements.items():
        if val is None:
            val = ''
        new = new.replace(f"{{{{{tag}}}}}", val)

    if new == full:
        return

    for run in para.runs:
        for child in list(run._element):
            if child.tag in (W_T, W_BR, W_TAB, W_CR):
                run._element.remove(child)

    first_r = para.runs[0]._element
    buf = []

    def flush_text():
        if not buf:
            return
        t = OxmlElement('w:t')
        t.text = "".join(buf)
        t.set(qn('xml:space'), 'preserve')
        first_r.append(t)
        buf.clear()

    for ch in new:
        if ch == BR:
            flush_text()
            first_r.append(OxmlElement('w:br'))
        elif ch == TAB:
            flush_text()
            first_r.append(OxmlElement('w:tab'))
        elif ch == '\n':
            flush_text()
            first_r.append(OxmlElement('w:br'))
        else:
            buf.append(ch)
    flush_text()


def _dossier_replace_doc(doc, replacements):
    """
    Replace placeholders in body paragraphs, table cells, and existing
    headers/footers.

    CRITICAL: do NOT touch section.header.paragraphs / section.footer.paragraphs
    unconditionally — python-docx materialises a default empty header/footer
    just by accessing those properties, which mutates the section XML and
    adds vertical space that pushes content onto a second page on every render.
    Check the underlying <w:headerReference>/<w:footerReference> XML first and
    only descend if a header/footer was actually defined in the template.
    """
    from docx.oxml.ns import qn

    H_REF = qn('w:headerReference')
    F_REF = qn('w:footerReference')

    for para in doc.paragraphs:
        _dossier_replace_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _dossier_replace_paragraph(para, replacements)

    for section in doc.sections:
        sectPr = section._sectPr
        if sectPr.find(H_REF) is not None:
            for para in section.header.paragraphs:
                _dossier_replace_paragraph(para, replacements)
        if sectPr.find(F_REF) is not None:
            for para in section.footer.paragraphs:
                _dossier_replace_paragraph(para, replacements)


def _dossier_prevent_table_break(doc):
    """
    Keep table rows from splitting across pages, WITHOUT clipping cell content.

    Correct behaviour:
      - Set cantSplit='1' so a row never breaks across pages.
      - Leave trHeight alone. 'atLeast' is the right rule: rows grow to fit
        their content, and Word pushes the whole row to the next page only if
        it physically can't fit on the current one.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for table in doc.tables:
        for row in table.rows:
            tr = row._tr
            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                tr.insert(0, trPr)

            for old in trPr.findall(qn('w:cantSplit')):
                trPr.remove(old)
            cant = OxmlElement('w:cantSplit')
            cant.set(qn('w:val'), '1')
            trPr.append(cant)

            # Intentionally do NOT modify <w:trHeight>.


def _dossier_build_replacements(req: DossierRequest) -> Dict[str, str]:
    today = req.today_date or datetime.now().strftime("%d/%m/%Y")
    date  = req.date or today
    # Friendly port label
    port_label = req.port.replace("-anch", " Anch.").replace("-", " ").title()

    # Build cargo and bl_weight strings from cargo_items if available
    if req.cargo_items:
        cargo_lines = [item.description for item in req.cargo_items if item.description]
        weight_lines = [item.weight for item in req.cargo_items if item.weight]
        cargo_str = "\n".join(cargo_lines) if cargo_lines else (req.cargo or "")
        weight_str = "\n".join(weight_lines) if weight_lines else (req.bl_weight or "")
    else:
        cargo_str = req.cargo or ""
        weight_str = req.bl_weight or ""

    return {
        "vessel_name":    req.vessel_name  or "",
        "imo":            req.imo          or "",
        "flag":           req.flag         or "",
        "loa":            req.loa          or "",
        "deadweight":     req.deadweight   or "",
        "gross_tonnage":  req.gross_tonnage or "",
        "owner":          req.owner        or "",
        "cargo":          cargo_str,
        "bl_weight":      weight_str,
        "shipper":        req.shipper      or "",
        "notify":         req.notify       or "",
        "from":           req.from_port    or "",
        "to":             req.to_port      or "",
        "bc":             req.bc           or "",
        "arrival_date":   req.arrival_date  or "",
        "berthing_date":  req.berthing_date or "",
        "departure_date": req.departure_date or "",
        "date":           date,
        "today_date":     today,
        "port":           port_label,
        "agent_count":    req.agent_count  or "",
        "ste_garde":      req.ste_garde    or "",
        "expimp":         req.expimp or (req.operation or "import").title(),
        "shift":          req.shift        or "",
    }


@app.post("/dossier/generate")
async def dossier_generate(req: DossierRequest, request: Request):
    """
    Download selected .docx templates from Cloudflare Pages,
    fill {{placeholders}}, bundle into a zip, return zip bytes.
    """
    if API_SECRET:
        client_secret = request.headers.get("X-API-Secret", "")
        if client_secret != API_SECRET:
            raise HTTPException(status_code=401, detail="Unauthorized")

    if not req.templates:
        raise HTTPException(status_code=400, detail="No templates selected")
    if not req.port:
        raise HTTPException(status_code=400, detail="Port is required")

    replacements = _dossier_build_replacements(req)
    logger.info(
        f"Dossier generate | IMO {req.imo} | port={req.port} "
        f"| templates={req.templates}"
    )

    zip_buffer  = io.BytesIO()
    errors: List[str] = []

    async with httpx.AsyncClient(timeout=20) as client:
        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for tpl_id in req.templates:
                filename = DOSSIER_TEMPLATE_FILES.get(tpl_id)
                if not filename:
                    errors.append(f"Unknown template: {tpl_id}")
                    continue

                url = f"{DOSSIER_PAGES_BASE}/{req.port}/{filename}"
                try:
                    r = await client.get(url)
                    r.raise_for_status()
                    docx_bytes = r.content
                except Exception as e:
                    logger.warning(f"Failed to fetch template {filename}: {e}")
                    errors.append(f"Could not fetch {filename}: {e}")
                    continue

                try:
                    doc = Document(io.BytesIO(docx_bytes))
                    _dossier_replace_doc(doc, replacements)
                    _dossier_prevent_table_break(doc)
                    out = io.BytesIO()
                    doc.save(out)
                    out.seek(0)
                    zf.writestr(filename, out.read())
                    logger.info(f"Filled {filename}")
                except Exception as e:
                    logger.error(f"Error filling {filename}: {e}", exc_info=True)
                    errors.append(f"Error filling {filename}: {e}")

    zip_buffer.seek(0)
    zip_bytes = zip_buffer.read()

    if not zip_bytes or len(zip_bytes) < 50:
        detail = "No documents generated."
        if errors:
            detail += " Errors: " + "; ".join(errors)
        raise HTTPException(status_code=500, detail=detail)

    if errors:
        logger.warning(f"Dossier completed with {len(errors)} error(s): {errors}")

    vessel_slug = re.sub(r"\s+", "_", (req.vessel_name or "VESSEL").upper())
    port_slug   = req.port.upper().replace("-", "_")
    date_slug   = datetime.now().strftime("%Y%m%d")
    out_filename = f"DOSSIER_{vessel_slug}_{port_slug}_{date_slug}.zip"

    return Response(
        content=zip_bytes,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{out_filename}"'},
    )
